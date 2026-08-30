import sys
import re
import json
from docx import Document

def parse_wrap_docx(filepath):
    doc = Document(filepath)
    
    data = {
        "report_date": "24 August 2026",
        "url_date": "2026-08-24",
        "report_date_long": "Monday, 24th August 2026",
        "headline": "",
        "introduction_paragraphs": [],
        "sections": [],
        "snapshot": {}
    }
    
    current_section = None
    parsing_intro = True
    
    table_index = 0
    tables = doc.tables
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        if not data["headline"] and not re.match(r'^\d+\.', text):
            if "Daily DSE Wrap" in text or "August 2026" in text:
                continue
            data["headline"] = text
            continue
            
        match = re.match(r'^(\d+)\.\s+(.*)', text)
        if match:
            parsing_intro = False
            if current_section:
                data["sections"].append(current_section)
            current_section = {
                "title": f"{match.group(1)}. {match.group(2)}",
                "paragraphs": [],
                "list": [],
                "has_table": False,
                "table": None
            }
            if table_index < len(tables) and match.group(1) in ["1", "2", "4"]:
                table = tables[table_index]
                parsed_table = {"headers": [], "rows": []}
                for i, row in enumerate(table.rows):
                    row_data = []
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip().replace('\n', ' ')
                        if i == 0:
                            parsed_table["headers"].append(cell_text)
                        else:
                            color = ""
                            if "+" in cell_text and not "% band" in cell_text:
                                color = "--gain"
                            elif "–" in cell_text or "-" in cell_text:
                                if "%" in cell_text or "pts" in cell_text:
                                    color = "--loss"
                            row_data.append({"text": cell_text, "color": color})
                    if i > 0 and any(r["text"] for r in row_data):
                        parsed_table["rows"].append(row_data)
                
                current_section["has_table"] = True
                current_section["table"] = parsed_table
                table_index += 1
            continue
            
        if text.startswith("•") or text.startswith("*"):
            if current_section:
                text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text.lstrip("•* ").strip())
                current_section["list"].append(text)
            continue
            
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = text.replace("’", "'")
        text = re.sub(r'^(Upper limit:)', r'<strong>\1</strong>', text)
        text = re.sub(r'^(Lower limit:)', r'<strong>\1</strong>', text)
        
        if parsing_intro:
            data["introduction_paragraphs"].append(text)
        else:
            if current_section:
                current_section["paragraphs"].append(text)

    if current_section:
        data["sections"].append(current_section)
        
    # Extract DSEI, TSI, Turnover from Section 1 table
    if data["sections"] and data["sections"][0]["has_table"]:
        tbl = data["sections"][0]["table"]
        for row in tbl["rows"]:
            if len(row) >= 2:
                label = row[0]["text"].lower()
                val = row[1]["text"]
                if "dsei" in label or "all-share" in label:
                    data["snapshot"]["dsei"] = val
                if "tsi" in label or "tanzania share" in label:
                    data["snapshot"]["tsi"] = val
                if "turnover" in label:
                    data["snapshot"]["turnover"] = val
                    
    return data

def parse_prices_docx(filepath):
    doc = Document(filepath)
    prices = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text.startswith('Current Prices') or text.startswith('Ticker'):
            continue
            
        parts = text.split('\t')
        if len(parts) >= 7:
            ticker = parts[0].strip()
            company = parts[1].strip()
            sector = parts[2].strip()
            last_price = parts[3].strip()
            change_pct = parts[4].strip()
            volume = parts[5].strip()
            turnover = parts[6].strip()
            
            prices.append({
                "ticker": ticker,
                "company": company,
                "sector": sector,
                "last_price_tzs": last_price,
                "change_pct": change_pct,
                "volume": volume,
                "turnover_tzs": turnover
            })
    return prices

if __name__ == "__main__":
    wrap_path = r"C:\Users\tmalu\Documents\Daily DSE Wrap 24 August 2026.docx"
    prices_path = r"C:\Users\tmalu\Documents\Current Prices 24 August 2026.docx"
    
    try:
        wrap_data = parse_wrap_docx(wrap_path)
        with open("parsed_wrap_24aug.json", "w", encoding="utf-8") as f:
            json.dump(wrap_data, f, indent=2)
        print("Successfully parsed Wrap.")
    except Exception as e:
        print(f"Error parsing wrap: {e}")
        
    try:
        prices_data = parse_prices_docx(prices_path)
        with open("parsed_prices_24aug.json", "w", encoding="utf-8") as f:
            json.dump(prices_data, f, indent=2)
        print("Successfully parsed Prices.")
    except Exception as e:
        print(f"Error parsing prices: {e}")
