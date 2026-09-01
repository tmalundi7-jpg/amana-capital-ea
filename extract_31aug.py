import os
from docx import Document
import json

def extract_docx(filepath):
    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}"}
    
    try:
        doc = Document(filepath)
        paras = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paras.append(t)
                
        tables = []
        for table in doc.tables:
            tbl_data = []
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                tbl_data.append(row_data)
            tables.append(tbl_data)
            
        return {"paragraphs": paras, "tables": tables}
    except Exception as e:
        return {"error": str(e)}

wrap_data = extract_docx(r"C:\Users\tmalu\Documents\Daily DSE Wrap 31 August 2026.docx")
prices_data = extract_docx(r"C:\Users\tmalu\Documents\Current Prices 31 August 2026.docx")

with open('extracted_31aug.json', 'w', encoding='utf-8') as f:
    json.dump({'wrap': wrap_data, 'prices': prices_data}, f, indent=2)
print("Data extracted to extracted_31aug.json")
