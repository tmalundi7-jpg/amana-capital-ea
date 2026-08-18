import sys
import re
import json
from docx import Document
from jinja2 import Environment, FileSystemLoader

def parse_docx(filepath):
    doc = Document(filepath)
    
    data = {
        "report_date": "17 August 2026",
        "url_date": "2026-08-17",
        "report_date_long": "Monday, 17th August 2026",
        "headline": "",
        "introduction_paragraphs": [],
        "sections": []
    }
    
    current_section = None
    parsing_intro = True
    
    table_index = 0
    tables = doc.tables
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if it's the headline (first non-empty paragraph after date)
        if not data["headline"] and not re.match(r'^\d+\.', text):
            # The date usually appears at the very top. If text is the date, skip it.
            if "Daily DSE Wrap" in text or "August 2026" in text:
                continue
            data["headline"] = text
            continue
            
        # Check for section header (e.g., "1. Market Snapshot")
        match = re.match(r'^(\d+)\.\s+(.*)', text)
        if match:
            parsing_intro = False
            if current_section:
                data["sections"].append(current_section)
            current_section = {
                "title": f"{match.group(1)}. {match.group(2)}",
                "paragraphs": [],
                "list": [],
                "has_table": False,
                "table": None
            }
            # Look ahead for a table that belongs to this section
            if table_index < len(tables) and match.group(1) in ["1", "2", "4"]:
                # Naive assumption: Sections 1, 2, and 4 typically have tables
                table = tables[table_index]
                parsed_table = {"headers": [], "rows": []}
                for i, row in enumerate(table.rows):
                    row_data = []
                    for j, cell in enumerate(row.cells):
                        # clean cell text
                        cell_text = cell.text.strip().replace('\n', ' ')
                        
                        if i == 0:
                            parsed_table["headers"].append(cell_text)
                        else:
                            # Try to infer color
                            color = ""
                            if "+" in cell_text and not "% band" in cell_text:
                                color = "--gain"
                            elif "–" in cell_text or "-" in cell_text:
                                if "%" in cell_text or "pts" in cell_text:
                                    color = "--loss"
                            
                            row_data.append({
                                "text": cell_text,
                                "color": color
                            })
                    if i > 0:
                        parsed_table["rows"].append(row_data)
                
                current_section["has_table"] = True
                current_section["table"] = parsed_table
                table_index += 1
            
            continue
            
        # If it's a list item (starts with a bullet or similar)
        if text.startswith("•") or text.startswith("*"):
            if current_section:
                # Format strong text if needed
                text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text.lstrip("•* ").strip())
                current_section["list"].append(text)
            continue
            
        # It's a paragraph
        # Basic markdown to HTML bolding
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        text = text.replace("’", "'")
        text = re.sub(r'^(Upper limit:)', r'<strong>\1</strong>', text)
        text = re.sub(r'^(Lower limit:)', r'<strong>\1</strong>', text)
        
        if parsing_intro:
            data["introduction_paragraphs"].append(text)
        else:
            if current_section:
                current_section["paragraphs"].append(text)

    if current_section:
        data["sections"].append(current_section)
        
    return data

def generate_html(data, template_path, output_path):
    env = Environment(loader=FileSystemLoader('.'))
    template = env.get_template(template_path)
    
    html_out = template.render(**data)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_out)
    print(f"Generated {output_path}")

if __name__ == "__main__":
    docx_path = r"C:\Users\tmalu\Documents\Daily DSE Wrap 17 August 2026.docx"
    template_path = "wrap_template.html"
    output_path = "dse-wrap-2026-08-17.html"
    
    data = parse_docx(docx_path)
    
    # Manually tweak for known structure if needed
    
    generate_html(data, template_path, output_path)
