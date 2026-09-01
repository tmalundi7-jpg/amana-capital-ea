import os
from docx import Document

def print_docx_content(filepath):
    print(f"--- CONTENT OF: {filepath} ---")
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    
    try:
        doc = Document(filepath)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"P{i}: {text}")
                
        print("\n--- TABLES ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}:")
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(" | ".join(row_data))
    except Exception as e:
        print(f"Error reading docx: {e}")

print_docx_content(r"C:\Users\tmalu\Documents\Daily DSE Wrap 31 August 2026.docx")
print_docx_content(r"C:\Users\tmalu\Documents\Current Prices 31 August 2026.docx")
